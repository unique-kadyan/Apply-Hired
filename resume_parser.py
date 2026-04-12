"""Resume parser — extracts profile data from PDF/DOCX/TXT files."""

import json
import logging
import os
import re
from pathlib import Path
from typing import Optional

from config import OPENAI_API_KEY

logger = logging.getLogger(__name__)


def _quality_score(text: str) -> int:
    """Rate how well-extracted a PDF text is. Higher = better."""
    if not text or not text.strip():
        return 0
    score = 0
    # Longer text is generally better (more content preserved)
    score += min(30, len(text) // 100)
    # Lines that look like complete job titles (multi-word, capitalized)
    title_patterns = len(re.findall(
        r'(?:Senior|Lead|Staff|Principal|Junior|Full|Software|Backend|Frontend|Project|Freelance)\s*[-–]?\s*\w+',
        text, re.IGNORECASE
    ))
    score += min(20, title_patterns * 5)
    # Bullet points preserved
    score += min(15, len(re.findall(r'[•\-\*]\s+\w', text)) * 2)
    # Section headers present
    headers = len(re.findall(r'\b(?:EXPERIENCE|EDUCATION|SKILLS|SUMMARY|PROJECTS|CERTIFICATIONS|ACHIEVEMENTS)\b', text, re.IGNORECASE))
    score += min(15, headers * 3)
    # Penalty for fragmented short lines (sign of bad extraction)
    lines = [ln.strip() for ln in text.split('\n') if ln.strip()]
    short_fragments = sum(1 for ln in lines if len(ln) < 4 and ln[0].isalpha())
    score -= short_fragments * 5
    # Penalty for broken words (e.g. "Sof" on one line, "tware" on next)
    broken = len(re.findall(r'\b[A-Z][a-z]{0,3}\n\s*[a-z]', text))
    score -= broken * 10
    return max(0, score)


def extract_text_from_pdf(filepath: str) -> str:
    """Extract text from a PDF by trying multiple strategies and picking the best."""
    candidates = []

    # Strategy 1: pdfplumber with layout mode
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            page_width = pdf.pages[0].width if pdf.pages else 600
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text(
                    layout=True,
                    layout_width=int(page_width / 4),
                    layout_width_chars=int(page_width / 4),
                )
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                candidates.append(("pdfplumber_layout", text))
    except Exception as e:
        logger.debug(f"pdfplumber layout: {e}")

    # Strategy 2: pdfplumber default (no layout)
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                candidates.append(("pdfplumber_default", text))
    except Exception as e:
        logger.debug(f"pdfplumber default: {e}")

    # Strategy 3: pdfplumber with generous tolerances
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=5, y_tolerance=5)
                if page_text:
                    text += page_text + "\n"
            if text.strip():
                candidates.append(("pdfplumber_tolerant", text))
    except Exception as e:
        logger.debug(f"pdfplumber tolerant: {e}")

    # Strategy 4: pdfminer (comes with pdfplumber)
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(filepath)
        if text and text.strip():
            candidates.append(("pdfminer", text))
    except Exception as e:
        logger.debug(f"pdfminer: {e}")

    # Strategy 5: PyPDF2 fallback
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            candidates.append(("pypdf2", text))
    except Exception as e:
        logger.debug(f"PyPDF2: {e}")

    if not candidates:
        return ""

    # Score each candidate and pick the best
    best_name, best_text = max(candidates, key=lambda c: _quality_score(c[1]))
    logger.info(f"PDF extraction: chose {best_name} (score={_quality_score(best_text)}) "
                f"from {len(candidates)} strategies")
    return best_text


def extract_text_from_docx(filepath: str) -> str:
    """Extract text from a DOCX file."""
    from docx import Document
    doc = Document(filepath)
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text_from_txt(filepath: str) -> str:
    """Extract text from a TXT file."""
    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(filepath: str) -> str:
    """Extract text from any supported file format."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".txt":
        return extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


# ---------------------------------------------------------------------------
# AI-powered parsing (multi-provider with automatic failover)
# ---------------------------------------------------------------------------


_RESUME_PARSE_PROMPT = """You are a professional resume parser. The text below was extracted from a PDF file using automated text extraction.

PDF text extraction commonly produces artifacts: words, job titles, company names, and sentences may be split across multiple lines or truncated mid-word due to column layouts, page breaks, or formatting. You must intelligently reconstruct all fragmented text into its complete, meaningful form by inferring from surrounding context. Never return partial or truncated words — every field must contain complete, properly formed text.

IMPORTANT skill classification rules:
- "Java" and "JavaScript" are DIFFERENT languages. Only list "Java" if the resume explicitly mentions Java (not as part of JavaScript). Only list "JavaScript" if explicitly mentioned.
- "React" and "React Native" are different. "React.js" = "React".
- "Node.js" is backend, not a language.
- "SQL" is a language. "PostgreSQL", "MySQL", "Oracle SQL" are databases.
- "Spring Boot", "Spring Security", "Spring Data" are separate backend skills — list each one mentioned.
- "AWS" is cloud. Individual services like "EC2", "Lambda", "S3" should also be listed under cloud_devops.
- Only list skills that are ACTUALLY mentioned in the resume. Do not infer or assume.

Parse the resume and return ONLY valid JSON with this structure:

{
    "name": "Full name",
    "title": "Current or most recent complete job title",
    "email": "email address",
    "phone": "phone number",
    "location": "City, Country",
    "years_of_experience": 0.0,
    "open_to": "role preferences if mentioned",
    "summary": "Professional summary, 2-4 sentences",
    "skills": {
        "languages": [],
        "backend": [],
        "frontend": [],
        "databases": [],
        "cloud_devops": [],
        "architecture": [],
        "testing": [],
        "cs_fundamentals": []
    },
    "experience": [
        {
            "title": "Complete job title",
            "company": "Complete company name",
            "period": "Start – End",
            "highlights": ["Complete bullet points"]
        }
    ],
    "education": "Degree | University | Score | Year",
    "certifications": [],
    "achievements": []
}

RESUME TEXT:
"""


def _pdf_to_base64_images(filepath: str) -> list[str]:
    """Convert each PDF page to a base64-encoded PNG image."""
    import base64
    images = []
    try:
        import pypdfium2 as pdfium
        pdf = pdfium.PdfDocument(filepath)
        for i in range(len(pdf)):
            page = pdf[i]
            # Render at 200 DPI for good quality without being too large
            bitmap = page.render(scale=200 / 72)
            pil_image = bitmap.to_pil()
            import io
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG", optimize=True)
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")
            images.append(b64)
            if i >= 2:  # max 3 pages
                break
        pdf.close()
    except Exception as e:
        logger.warning(f"PDF to image conversion failed: {e}")
    return images


def _build_ai_providers() -> list[dict]:
    """Build an ordered list of AI providers to try.

    Strategy: free providers FIRST, paid providers LAST.
    All use OpenAI-compatible /v1/chat/completions format.
    """
    providers = []

    # --- FREE providers (confirmed working, ordered by speed/quality) ---

    # 1. Groq (free, fastest inference, 70B model)
    groq_key = os.environ.get("GROQ_API_KEY", "")
    if groq_key:
        providers.append({
            "name": "groq",
            "base_url": "https://api.groq.com/openai/v1",
            "api_key": groq_key,
            "model": "llama-3.3-70b-versatile",
            "supports_vision": False,
        })

    # 2. Cerebras (free, fast, 30 req/min)
    cerebras_key = os.environ.get("CEREBRAS_API_KEY", "")
    if cerebras_key:
        providers.append({
            "name": "cerebras",
            "base_url": "https://api.cerebras.ai/v1",
            "api_key": cerebras_key,
            "model": "llama3.1-8b",
            "supports_vision": False,
        })

    # 3. Mistral (free, 1B tokens/month)
    mistral_key = os.environ.get("MISTRAL_API_KEY", "")
    if mistral_key:
        providers.append({
            "name": "mistral",
            "base_url": "https://api.mistral.ai/v1",
            "api_key": mistral_key,
            "model": "mistral-small-latest",
            "supports_vision": False,
        })

    # 4. SambaNova (free, 70B model)
    sambanova_key = os.environ.get("SAMBANOVA_API_KEY", "")
    if sambanova_key:
        providers.append({
            "name": "sambanova",
            "base_url": "https://api.sambanova.ai/v1",
            "api_key": sambanova_key,
            "model": "Meta-Llama-3.3-70B-Instruct",
            "supports_vision": False,
        })

    # 5. Novita AI (free models)
    novita_key = os.environ.get("NOVITA_API_KEY", "")
    if novita_key:
        providers.append({
            "name": "novita",
            "base_url": "https://api.novita.ai/v3/openai",
            "api_key": novita_key,
            "model": "qwen/qwen-2.5-72b-instruct",
            "supports_vision": False,
        })

    # 6. Pollinations (no key needed)
    providers.append({
        "name": "pollinations",
        "base_url": "https://text.pollinations.ai/openai",
        "api_key": "dummy",
        "model": "openai",
        "supports_vision": False,
    })

    # 7. GPT-OSS (no key needed)
    gptoss_url = os.environ.get("GPTOSS_URL", "https://broken-water-d859.junioralive.workers.dev/v1")
    providers.append({
        "name": "gpt-oss",
        "base_url": gptoss_url,
        "api_key": "dummy",
        "model": "gpt-oss-20b",
        "supports_vision": False,
    })

    # 8. OllamaFreeAPI (no key, community nodes)
    providers.append({
        "name": "ollama-free",
        "base_url": None,
        "api_key": "dummy",
        "model": "llama3.2:3b",
        "supports_vision": False,
        "custom": True,
    })

    # --- RATE-LIMITED free (may fail temporarily) ---

    # 9. Google Gemini (free but rate limited on new keys)
    gemini_key = os.environ.get("GEMINI_API_KEY", "")
    if gemini_key:
        providers.append({
            "name": "gemini",
            "base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
            "api_key": gemini_key,
            "model": "gemini-2.0-flash",
            "supports_vision": True,
        })

    # --- PAID (last resort) ---

    if OPENAI_API_KEY:
        providers.append({
            "name": "openai",
            "base_url": None,
            "api_key": OPENAI_API_KEY,
            "model": "gpt-4o-mini",
            "supports_vision": True,
        })

    return providers


def _call_ai_vision(provider: dict, prompt: str, images: list[str]) -> Optional[str]:
    """Call an AI provider with images (vision mode). Returns raw response text."""
    try:
        from openai import OpenAI
    except ImportError:
        return None

    kwargs = {"api_key": provider["api_key"]}
    if provider["base_url"]:
        kwargs["base_url"] = provider["base_url"]

    # Build multimodal content: text prompt + images
    content = [{"type": "text", "text": prompt}]
    for img_b64 in images:
        content.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{img_b64}", "detail": "high"},
        })

    client = OpenAI(**kwargs)
    response = client.chat.completions.create(
        model=provider["model"],
        max_tokens=4000,
        messages=[{"role": "user", "content": content}],
    )
    return response.choices[0].message.content


def _call_ai_text(provider: dict, prompt: str) -> Optional[str]:
    """Call an AI provider with text-only prompt."""
    # Custom handler for OllamaFreeAPI (non-OpenAI SDK)
    if provider.get("custom") and provider["name"] == "ollama-free":
        try:
            from ollamafreeapi import OllamaFreeAPI
            client = OllamaFreeAPI()
            return client.chat(model=provider["model"], prompt=prompt)
        except Exception:
            return None

    try:
        from openai import OpenAI
    except ImportError:
        return None

    kwargs = {"api_key": provider["api_key"]}
    if provider["base_url"]:
        kwargs["base_url"] = provider["base_url"]

    client = OpenAI(**kwargs)
    response = client.chat.completions.create(
        model=provider["model"],
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}],
    )
    return response.choices[0].message.content


def _is_quota_error(e: Exception) -> bool:
    """Check if an exception is a quota/rate-limit error."""
    err = str(e).lower()
    return any(kw in err for kw in ("quota", "rate_limit", "429", "insufficient_quota", "billing"))


def _parse_ai_response(result_text: str) -> Optional[dict]:
    """Extract JSON from AI response text."""
    if not result_text:
        return None
    json_match = re.search(r"\{.*\}", result_text, re.DOTALL)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError:
            pass
    return None


def parse_resume_ai(text: str, filepath: str = None) -> Optional[dict]:
    """Parse resume using AI with automatic failover.

    Strategy:
    1. Vision providers get PDF page images (zero text extraction artifacts)
    2. Text providers get extracted text as fallback
    3. Each provider is tried in priority order
    """
    providers = _build_ai_providers()
    if not providers:
        return None

    # Convert PDF to images for vision-capable providers
    images = []
    if filepath and filepath.lower().endswith(".pdf"):
        images = _pdf_to_base64_images(filepath)

    vision_prompt = _RESUME_PARSE_PROMPT.replace(
        "The text below was extracted from a PDF file using automated text extraction.\n\n"
        "PDF text extraction commonly produces artifacts: words, job titles, company names, "
        "and sentences may be split across multiple lines or truncated mid-word due to column "
        "layouts, page breaks, or formatting. You must intelligently reconstruct all fragmented "
        "text into its complete, meaningful form by inferring from surrounding context. Never "
        "return partial or truncated words — every field must contain complete, properly formed text.",
        "Parse the resume shown in the image(s) below. Extract ALL information exactly as displayed — "
        "every job title, company, date, bullet point, skill, and section. Do not skip or truncate anything."
    )

    text_prompt = _RESUME_PARSE_PROMPT + text[:12000]

    for provider in providers:
        try:
            logger.info(f"Trying resume parse with {provider['name']} "
                        f"({'vision' if provider.get('supports_vision') and images else 'text'})")

            # Use vision if provider supports it and we have images
            if provider.get("supports_vision") and images:
                result_text = _call_ai_vision(provider, vision_prompt, images)
            else:
                result_text = _call_ai_text(provider, text_prompt)

            parsed = _parse_ai_response(result_text)
            if parsed:
                logger.info(f"Resume parsed with {provider['name']}")
                return parsed

        except Exception as e:
            if _is_quota_error(e):
                logger.warning(f"{provider['name']} quota exceeded, trying next provider")
                continue
            logger.error(f"{provider['name']} parsing failed: {e}")
            continue

    return None


# ---------------------------------------------------------------------------
# Regex-based parsing (no API key needed)
# ---------------------------------------------------------------------------

# Common skill keywords to look for
# Each entry is (regex_pattern, display_name). If display_name is None, the pattern is used as-is.
SKILL_PATTERNS = {
    "languages": [
        ("Java(?!Script)", "Java"), "Python", "JavaScript", "TypeScript", ("C\\+\\+", "C++"), "C#", "Go",
        "Golang", "Rust", "Ruby", "PHP", "Kotlin", "Swift", "Scala", "SQL",
        "Perl", "Dart", "Lua", "Shell", "Bash",
    ],
    "backend": [
        "Spring Boot", "Spring Security", "Spring Data", "Spring Cloud",
        "Hibernate", "FastAPI", "Django", "Flask", ("Express\\.js", "Express.js"), "NestJS",
        ("Node\\.js", "Node.js"), "Kafka", "RabbitMQ", "gRPC", "GraphQL", "REST API",
        ("ASP\\.NET", "ASP.NET"), "Ruby on Rails", "Laravel", "Gin", "Fiber",
    ],
    "frontend": [
        "React", ("React\\.js", "React.js"), ("Next\\.js", "Next.js"), ("Vue\\.js", "Vue.js"), "Angular",
        "Svelte", "Tailwind", "Bootstrap", ("HTML5?", "HTML5"), ("CSS3?", "CSS3"),
        "jQuery", "Redux", "Webpack", "Vite",
    ],
    "databases": [
        "PostgreSQL", "MySQL", "MongoDB", "Redis", ("Oracle SQL", "Oracle SQL"),
        ("Oracle DB", "Oracle DB"), ("Oracle", "Oracle SQL"),
        "SQLite", "DynamoDB", "Cassandra", "Elasticsearch",
        "Firebase", "CouchDB", "MariaDB", "SQL Server", "Neo4j",
    ],
    "cloud_devops": [
        "AWS", "GCP", "Azure", "Docker", "Kubernetes", "K8s",
        "Terraform", "Ansible", "Jenkins", "GitHub Actions", "GitLab CI",
        "CI/CD", "CloudWatch", "EC2", "Lambda", "S3", "RDS", "SQS",
        "ECS", "EKS", "Fargate", "Heroku", "Vercel", "Netlify",
    ],
    "architecture": [
        "Microservices", "REST APIs", ("Event[- ]Driven", "Event-Driven"), ("Distributed Systems?", "Distributed Systems"),
        "HLD", "LLD", "DDD", ("Domain[- ]Driven", "Domain-Driven"), "CQRS", "Serverless",
        "SOA", "Monolith", "API Gateway",
    ],
    "testing": [
        "JUnit", "Mockito", "Selenium", "Cypress", "Jest", "Mocha",
        "PyTest", "TDD", "BDD", "Integration Testing", "Unit Testing",
        "Performance Testing", "Load Testing",
    ],
}


def _extract_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group() if match else ""


def _extract_phone(text: str) -> str:
    match = re.search(r"[\+]?[\d\s\-\(\)]{10,15}", text)
    return match.group().strip() if match else ""


def _extract_name(text: str) -> str:
    """Attempt to extract name from first few lines."""
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    for line in lines[:5]:
        # Name is typically the first non-empty line, all caps or title case
        clean = re.sub(r"[^a-zA-Z\s]", "", line).strip()
        words = clean.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return clean
    return lines[0] if lines else ""


def _extract_skills(text: str) -> dict[str, list[str]]:
    """Extract skills by matching known patterns."""
    found = {}
    for group, entries in SKILL_PATTERNS.items():
        matches = []
        for entry in entries:
            if isinstance(entry, tuple):
                pattern, display = entry
            else:
                pattern, display = entry, entry
            if re.search(rf"\b{pattern}\b", text, re.IGNORECASE):
                if display not in matches:
                    matches.append(display)
        found[group] = matches
    return found


def _extract_experience_years(text: str) -> float:
    """Estimate years of experience from text."""
    match = re.search(r"(\d+(?:\.\d+)?)\s*(?:\+\s*)?years?\s*(?:of\s*)?(?:experience)?", text, re.IGNORECASE)
    if match:
        return float(match.group(1))

    # Try to calculate from date ranges
    year_pattern = re.findall(r"(20\d{2})\s*[–\-]\s*(20\d{2}|Present)", text, re.IGNORECASE)
    if year_pattern:
        import datetime
        total_years = 0
        current_year = datetime.datetime.now().year
        for start, end in year_pattern:
            end_year = current_year if end.lower() == "present" else int(end)
            total_years += end_year - int(start)
        return round(total_years, 1)

    return 0.0


def _extract_summary(text: str) -> str:
    """Extract summary/objective section."""
    patterns = [
        r"(?:SUMMARY|OBJECTIVE|ABOUT|PROFILE)\s*[:\n]\s*(.+?)(?:\n\n|\n[A-Z]{2,})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            summary = match.group(1).strip()
            return re.sub(r"\s+", " ", summary)[:500]
    return ""


def _extract_experience(text: str) -> list[dict]:
    """Extract work experience entries."""
    entries = []
    # Look for patterns like "Title | Company" or "Title at Company" followed by dates
    exp_pattern = re.findall(
        r"(.+?)\s*[|–\-@at]+\s*(.+?)\s+((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*[–\-]\s*(?:Present|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}))",
        text,
        re.IGNORECASE,
    )

    for title, company, period in exp_pattern[:6]:
        title = title.strip().split("\n")[-1].strip()
        company = company.strip().split("\n")[0].strip()
        entries.append({
            "title": title[:100],
            "company": company[:100],
            "period": period.strip(),
            "highlights": [],
        })

    # Extract bullet points near each entry
    bullets = re.findall(r"[•\-\*]\s*(.+?)(?:\n|$)", text)
    if entries and bullets:
        per_entry = max(1, len(bullets) // max(len(entries), 1))
        for i, entry in enumerate(entries):
            start = i * per_entry
            entry["highlights"] = [b.strip()[:200] for b in bullets[start:start + per_entry]][:4]

    return entries


def _extract_education(text: str) -> str:
    """Extract education section."""
    match = re.search(
        r"(?:EDUCATION|ACADEMIC|QUALIFICATION)\s*[:\n]\s*(.+?)(?:\n\n|\n[A-Z]{2,}|$)",
        text, re.IGNORECASE | re.DOTALL,
    )
    if match:
        return re.sub(r"\s+", " ", match.group(1).strip())[:300]
    return ""


def _extract_certifications(text: str) -> list[str]:
    """Extract certifications."""
    match = re.search(
        r"(?:CERTIFICATION|CERTIFICATE|LICENSES?)\s*[:\n]\s*(.+?)(?:\n\n|\n[A-Z]{2,}|$)",
        text, re.IGNORECASE | re.DOTALL,
    )
    if match:
        certs = re.findall(r"[•\-\*]\s*(.+?)(?:\n|$)", match.group(1))
        if certs:
            return [c.strip()[:200] for c in certs][:5]
        return [match.group(1).strip()[:200]]
    return []


def parse_resume_local(text: str) -> dict:
    """Parse resume using regex patterns (no API key needed)."""
    return {
        "name": _extract_name(text),
        "title": "",
        "email": _extract_email(text),
        "phone": _extract_phone(text),
        "location": "",
        "years_of_experience": _extract_experience_years(text),
        "open_to": "Remote",
        "summary": _extract_summary(text),
        "skills": _extract_skills(text),
        "experience": _extract_experience(text),
        "education": _extract_education(text),
        "certifications": _extract_certifications(text),
    }


# ---------------------------------------------------------------------------
# Main parser
# ---------------------------------------------------------------------------

def _fix_broken_fields(parsed: dict) -> dict:
    """Post-process parsed resume to fix truncated titles and companies.

    Detects cases where the title is a word fragment and the company field
    contains the rest of the title merged with the actual company name.
    E.g. title="Sof", company="ware Engineer | IVY Comptech"
      → title="Software Engineer", company="IVY Comptech"
    """
    for exp in parsed.get("experience", []):
        title = (exp.get("title") or "").strip()
        company = (exp.get("company") or "").strip()

        if not title or not company:
            continue

        # Heuristic: title is broken if it's very short AND the company field
        # starts with lowercase or continues the word fragment
        title_looks_broken = (
            len(title) <= 8 and
            not title.endswith(("er", "or", "ist", "ant", "tic")) and
            (company[0].islower() or "|" in company)
        )

        if title_looks_broken and "|" in company:
            # company = "ware Engineer | IVY Comptech"  or  "Stack Engineer (...) | US, UK..."
            parts = company.split("|", 1)
            rest_of_title = parts[0].strip()
            actual_company = parts[1].strip() if len(parts) > 1 else ""

            merged = title + rest_of_title
            # Check if merging produces a real word (no space between fragment and continuation)
            # e.g. "Sof" + "tware Engineer" = "Software Engineer"
            if rest_of_title and rest_of_title[0].islower():
                exp["title"] = merged
            else:
                # e.g. "Full" + " Stack Engineer" — needs hyphen or space
                exp["title"] = (title + "-" + rest_of_title).replace("- ", "-") if title.endswith(("Full", "Co")) else title + " " + rest_of_title
                exp["title"] = exp["title"].strip()

            exp["company"] = actual_company

        elif title_looks_broken and company[0].islower():
            # No pipe — company starts with lowercase continuing the title
            # e.g. title="Projec", company="t Engineer at Wipro"
            merged = title + company
            # Try to split on known separators
            for sep in (" at ", " | ", " - "):
                if sep in merged:
                    parts = merged.split(sep, 1)
                    exp["title"] = parts[0].strip()
                    exp["company"] = parts[1].strip()
                    break
            else:
                exp["title"] = merged

    # Fix common misspelled words from PDF fragment merging
    _correct_titles(parsed)

    # Clean up garbage highlights (single numbers, very short fragments, tech tags)
    for exp in parsed.get("experience", []):
        if "highlights" in exp:
            exp["highlights"] = [
                h for h in exp["highlights"]
                if len(h) > 5 and not re.match(r'^\d{1,2}$', h.strip())
                and not re.match(r'^[A-Z][a-z]{0,2}$', h.strip())
            ]

    # Also fix the top-level title if it matches the first experience
    experience = parsed.get("experience", [])
    top_title = (parsed.get("title") or "").strip()
    if experience and top_title and len(top_title) <= 8:
        parsed["title"] = experience[0].get("title", top_title)

    return parsed


def _correct_titles(parsed: dict):
    """Fix misspelled words that result from PDF fragment merging."""
    # Common job title words — used for fuzzy correction
    known_words = {
        "software", "senior", "junior", "engineer", "developer", "architect",
        "manager", "director", "lead", "staff", "principal", "freelance",
        "full-stack", "fullstack", "frontend", "backend", "project", "product",
        "consultant", "analyst", "specialist", "administrator", "coordinator",
        "associate", "intern", "trainee", "applications", "platforms",
        "solutions", "technology", "technologies", "systems", "infrastructure",
        "devops", "security", "data", "cloud", "mobile", "remote",
    }

    def _fix_word(word: str) -> str:
        wl = word.lower()
        if wl in known_words:
            return word
        # Try inserting one letter at each position to match a known word
        for i in range(len(wl) + 1):
            for c in "abcdefghijklmnopqrstuvwxyz":
                candidate = wl[:i] + c + wl[i:]
                if candidate in known_words:
                    # Preserve original casing of first char
                    return (word[0] + candidate[1:]) if word[0].isupper() else candidate
        return word

    def _fix_text(text: str) -> str:
        return " ".join(_fix_word(w) if w[0].isalpha() else w for w in text.split())

    for exp in parsed.get("experience", []):
        exp["title"] = _fix_text(exp.get("title", ""))
        exp["company"] = _fix_text(exp.get("company", ""))


def parse_resume(filepath: str) -> dict:
    """Parse a resume file and return structured profile data.
    Uses AI if available, falls back to regex parsing."""
    text = extract_text(filepath)

    if not text.strip():
        raise ValueError("Could not extract text from the resume file.")

    # Try AI parsing (auto-failover: OpenAI vision → DeepSeek text → Pollinations text)
    ai_result = parse_resume_ai(text, filepath=filepath)
    if ai_result:
        logger.info("Resume parsed with AI")
        return _fix_broken_fields(ai_result)

    # Fall back to local regex parsing
    logger.info("Resume parsed with regex (no API key)")
    return _fix_broken_fields(parse_resume_local(text))


# ---------------------------------------------------------------------------
# Resume scorer
# ---------------------------------------------------------------------------

def _score_resume_local(text: str) -> dict:
    """Score a resume using heuristics (no API key needed)."""
    scores = {}
    total = 0

    # 1. Contact info (10 pts)
    contact = 0
    if _extract_email(text):
        contact += 3
    if _extract_phone(text):
        contact += 3
    if re.search(r"linkedin|github", text, re.IGNORECASE):
        contact += 4
    scores["contact_info"] = {"score": contact, "max": 10, "tips": []}
    if contact < 10:
        if not _extract_email(text):
            scores["contact_info"]["tips"].append("Add your email address")
        if not _extract_phone(text):
            scores["contact_info"]["tips"].append("Add your phone number")
        if not re.search(r"linkedin", text, re.IGNORECASE):
            scores["contact_info"]["tips"].append("Add your LinkedIn profile URL")
        if not re.search(r"github", text, re.IGNORECASE):
            scores["contact_info"]["tips"].append("Add your GitHub profile URL")
    total += contact

    # 2. Summary/Objective (10 pts)
    summary = _extract_summary(text)
    summary_score = 0
    if summary:
        words = len(summary.split())
        if words >= 30:
            summary_score = 10
        elif words >= 15:
            summary_score = 7
        else:
            summary_score = 4
    scores["summary"] = {"score": summary_score, "max": 10, "tips": []}
    if summary_score < 10:
        if not summary:
            scores["summary"]["tips"].append("Add a professional summary at the top of your resume")
        elif len(summary.split()) < 30:
            scores["summary"]["tips"].append("Expand your summary to 2-3 sentences with measurable impact")
    total += summary_score

    # 3. Skills (15 pts)
    skills = _extract_skills(text)
    skill_count = sum(len(v) for v in skills.values())
    skill_score = min(15, skill_count * 1.5)
    scores["skills"] = {"score": round(skill_score), "max": 15, "tips": []}
    if skill_score < 15:
        scores["skills"]["tips"].append(f"Found {skill_count} skills — aim for 10+ relevant technical skills")
        if not skills.get("cloud_devops"):
            scores["skills"]["tips"].append("Add cloud/DevOps skills (AWS, Docker, CI/CD)")
    total += skill_score

    # 4. Experience (25 pts)
    exp = _extract_experience(text)
    exp_score = 0
    bullet_count = 0
    metrics = 0
    if exp:
        exp_score += min(10, len(exp) * 3)  # up to 10 for number of roles
        bullet_count = sum(len(e.get("highlights", [])) for e in exp)
        exp_score += min(10, bullet_count * 2)  # up to 10 for bullets
        # Check for quantified achievements
        metrics = len(re.findall(r'\d+[%xX]|\$[\d,]+|\d+\+?\s*(?:users|clients|engineers|teams|rps)', text, re.IGNORECASE))
        exp_score += min(5, metrics * 1.5)  # up to 5 for metrics
    scores["experience"] = {"score": round(min(25, exp_score)), "max": 25, "tips": []}
    if exp_score < 25:
        if not exp:
            scores["experience"]["tips"].append("Add your work experience with company names and dates")
        elif bullet_count < 6:
            scores["experience"]["tips"].append("Add more bullet points to each role (3-5 per job)")
        if metrics < 3:
            scores["experience"]["tips"].append("Quantify achievements with numbers (%, $, users, etc.)")
    total += min(25, exp_score)

    # 5. Education (10 pts)
    edu = _extract_education(text)
    edu_score = 10 if edu and len(edu) > 10 else 0
    scores["education"] = {"score": edu_score, "max": 10, "tips": []}
    if edu_score == 0:
        scores["education"]["tips"].append("Add your education with degree, university, and year")
    total += edu_score

    # 6. Formatting & length (15 pts)
    word_count = len(text.split())
    format_score = 0
    if 300 <= word_count <= 1200:
        format_score += 5
    elif word_count > 100:
        format_score += 3
    if len(re.findall(r'[•\-\*]\s', text)) >= 5:
        format_score += 5  # uses bullet points
    if re.search(r'(?:EXPERIENCE|EDUCATION|SKILLS|SUMMARY|PROJECTS)', text, re.IGNORECASE):
        format_score += 5  # has section headers
    scores["formatting"] = {"score": min(15, format_score), "max": 15, "tips": []}
    if format_score < 15:
        if word_count < 300:
            scores["formatting"]["tips"].append("Resume seems too short — aim for 400-800 words")
        elif word_count > 1200:
            scores["formatting"]["tips"].append("Resume is too long — keep it to 1-2 pages")
        if len(re.findall(r'[•\-\*]\s', text)) < 5:
            scores["formatting"]["tips"].append("Use bullet points for achievements instead of paragraphs")
    total += min(15, format_score)

    # 7. Keywords & ATS (15 pts)
    ats_keywords = ["team", "lead", "manage", "develop", "design", "implement", "optimize",
                    "deploy", "scale", "architect", "mentor", "collaborate", "deliver", "automate"]
    keyword_hits = sum(1 for kw in ats_keywords if re.search(rf"\b{kw}", text, re.IGNORECASE))
    ats_score = min(15, keyword_hits * 2)
    scores["ats_keywords"] = {"score": round(ats_score), "max": 15, "tips": []}
    if ats_score < 15:
        scores["ats_keywords"]["tips"].append("Use action verbs: led, built, optimized, delivered, scaled, automated")
    total += ats_score

    return {
        "total_score": round(min(100, total)),
        "max_score": 100,
        "sections": scores,
    }


def _score_resume_ai(text: str, target_role: str = "") -> Optional[dict]:
    """Score a resume using AI providers with automatic failover.
    Passes target_role context so the AI evaluates ATS fit against the specific role."""
    providers = _build_ai_providers()
    if not providers:
        return None

    role_context = f"\nTARGET ROLE: {target_role}" if target_role else ""
    prompt = f"""You are an expert ATS resume scorer. Score this resume against the EXACT criteria below.
Award the FULL points for each section as long as the threshold is met — do NOT deduct for stylistic opinions.{role_context}

SCORING RUBRIC (award full marks when the threshold is met — no subjective deductions):

contact_info (max 10):
  10/10 → email AND phone AND LinkedIn URL AND GitHub/portfolio URL are all present
  7/10  → any 3 of the 4 present
  5/10  → any 2 of the 4 present
  2/10  → only 1 present

summary (max 10):
  10/10 → 4+ sentences AND 80+ words AND 3+ hard quantified metrics (%, x faster, $, ms, K+ users, etc.)
  7/10  → 3+ sentences AND 50+ words AND 1-2 metrics
  4/10  → short paragraph, no metrics

skills (max 15):
  15/15 → 20+ specific named technologies organized into categories
  10/15 → 12-19 named technologies
  6/15  → fewer than 12 technologies

experience (max 25):
  25/25 → every role has 5 bullet points AND every bullet starts with an action verb AND contains a hard metric
  18/25 → most roles have 4-5 bullets with some metrics
  10/25 → roles listed with 2-3 bullets, few or no metrics

education (max 10):
  10/10 → full degree name AND institution name AND graduation year all present
  6/10  → degree and institution present, year missing
  3/10  → institution only or partial info

formatting (max 15):
  15/15 → clear section headers (SUMMARY / SKILLS / EXPERIENCE / EDUCATION), consistent bullet style, no orphan lines
  10/15 → most sections present, minor inconsistencies
  5/15  → missing sections or inconsistent formatting

ats_keywords (max 15):
  15/15 → 20+ role-specific ATS keywords naturally woven into summary, bullets, and skills{"for the target role" if target_role else ""}
  10/15 → 10-19 keywords present
  5/15  → fewer than 10 role-relevant keywords

Return ONLY valid JSON — no markdown, no explanation:
{{
    "total_score": <integer 0-100>,
    "max_score": 100,
    "sections": {{
        "contact_info": {{ "score": <0-10>, "max": 10, "tips": [<specific actionable tips if score < 10, else []>] }},
        "summary": {{ "score": <0-10>, "max": 10, "tips": [] }},
        "skills": {{ "score": <0-15>, "max": 15, "tips": [] }},
        "experience": {{ "score": <0-25>, "max": 25, "tips": [] }},
        "education": {{ "score": <0-10>, "max": 10, "tips": [] }},
        "formatting": {{ "score": <0-15>, "max": 15, "tips": [] }},
        "ats_keywords": {{ "score": <0-15>, "max": 15, "tips": [] }}
    }}
}}

RESUME TEXT:
{text[:8000]}"""

    for provider in providers:
        try:
            logger.info(f"Scoring resume with {provider['name']}")
            result_text = _call_ai_text(provider, prompt)
            parsed = _parse_ai_response(result_text)
            if parsed and "total_score" in parsed:
                logger.info(f"Resume scored with {provider['name']}: {parsed.get('total_score')}/100")
                return parsed
        except Exception as e:
            if _is_quota_error(e):
                logger.warning(f"{provider['name']} quota exceeded, trying next")
                continue
            logger.error(f"{provider['name']} scoring failed: {e}")
            continue

    return None


def score_resume(filepath: str, target_role: str = "") -> dict:
    """Score a resume file. Uses AI providers with failover, falls back to heuristics."""
    text = extract_text(filepath)
    if not text.strip():
        raise ValueError("Could not extract text from the resume file.")

    ai_score = _score_resume_ai(text, target_role=target_role)
    if ai_score:
        ai_score["method"] = "ai"
        return ai_score

    result = _score_resume_local(text)
    result["method"] = "local"
    return result
